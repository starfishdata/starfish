# DOCX parsers
import os
from starfish.data_ingest.parsers.base_parser import BaseParser
from typing import Dict, Any


class WordDocumentParser(BaseParser):
    """Parser for Microsoft Word documents"""

    def __init__(self):
        super().__init__()
        self._docx = None
        self.supported_extensions = [".docx"]
        self.metadata = {}

    def _load_docx(self):
        """Lazy load the docx module"""
        if self._docx is None:
            try:
                import docx

                self._docx = docx
            except ImportError:
                raise ImportError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")

    def parse(self, file_path: str) -> str:
        """Parse a DOCX file into plain text

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text from the document
        """
        self._load_docx()
        doc = self._docx.Document(file_path)

        # Extract metadata
        self.metadata = {
            "author": doc.core_properties.author,
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified,
            "title": doc.core_properties.title,
            "pages": len(doc.paragraphs) // 50,  # Estimate pages
        }

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs]

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        return "\n\n".join(p for p in paragraphs if p)

    def get_metadata(self) -> Dict[str, Any]:
        """Get document metadata

        Returns:
            Dictionary containing document metadata
        """
        return self.metadata

    def is_supported(self, file_path: str) -> bool:
        """Check if the file is supported by this parser

        Args:
            file_path: Path to the file

        Returns:
            True if the file is supported, False otherwise
        """
        return os.path.splitext(file_path)[1].lower() in self.supported_extensions
